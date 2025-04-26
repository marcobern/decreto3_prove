from docx import Document
from docx.shared import Inches
import pandas as pd
import os

def add_dataframe_to_doc(doc, df, title="Table"):
    """Adds a DataFrame as a table to a Word document."""
    doc.add_paragraph(title, style='Heading 2')
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1]+1)
    table.style = 'Table Grid'

    # Add headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ""
    for i in range(df.shape[1]):
        hdr_cells[i + 1].text = str(i)

    # Add rows
    for i, (index, row) in enumerate(df.iterrows()):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = str(index)
        for i, value in enumerate(row):
            row_cells[i + 1].text = str(value)

    doc.add_paragraph("\n")  # Space after table


def add_image_to_doc(doc, image_path, caption="Image"):
    """Adds an image to the Word document."""
    doc.add_paragraph(caption, style='Heading 2')
    doc.add_picture(image_path, width=Inches(5))
    doc.add_paragraph("\n")  # Space after image


def generate_docx(entities, output_path="report.docx"):
    """Generates a Word document with tables and images."""
    doc = Document()
    doc.add_paragraph("Generated Report", style='Title')

    # Add dataframes or image
    for i, (name, source) in enumerate(entities):
        if isinstance(source, pd.DataFrame):
            add_dataframe_to_doc(doc, source, title=f"Table {i+1}: {name}")
        elif isinstance(source, str) and source.endswith((".png", ".jpg", ".jpeg")) and os.path.exists(source):
            add_image_to_doc(doc, source, caption=f"Figure {i+1}: {name}")

    # Save the document
    doc.save(output_path)
    print(f"Report saved as {output_path}")


if __name__ == "__main__":
    import numpy as np
    import matplotlib.pyplot as plt
    # Example usage:
    df1 = pd.DataFrame({'A': [1, 2, 3], 'B': [4, 5, 6]})
    df2 = pd.DataFrame({'X': ['a', 'b', 'c'], 'Y': ['d', 'e', 'f']})
    # Replace with actual image paths
    image_files = ["image1.png", "image2.jpg"]
    for img in image_files:
        plt.imshow(np.random.rand(100, 100))
        plt.savefig(img)

    generate_docx([df1, df2], image_files, "output_report.docx")

    for img in image_files:
        os.remove(img)
